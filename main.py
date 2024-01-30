from docx import Document
import re
import sys
import random
import tkinter as tk
from tkinter import simpledialog, messagebox, filedialog
from docx.shared import Pt
from docx.oxml.ns import qn

chinese_numbers = "一二三四五六七八九十"


# 将题库生成题库字典
def generate_bank_dict():
    pattern_type = re.compile(f'^[{chinese_numbers}]+、')  # 题型匹配法则
    pattern_question = re.compile(r'^(\d+)\.')  # 题目匹配法则
    question_type = ""  # 题型（顿号后面的）
    bank_dict = {}  # 键是题型，值是一个题目列表，列表包含一个个的题目字典
    question_list = []  # 题目列表
    question_dict = {}  # 题目字典
    question = ""  # 题目
    answer = ""  # 答案

    # 提示用户选择文件
    bank_path = filedialog.askopenfilename(
        title="选择题库文件",
        filetypes=[("Word files", "*.docx")])
    if not bank_path:
        sys.exit()
    tk = Document(bank_path)
    for paragraph in tk.paragraphs:
        # 记录答案
        ans = paragraph.text.strip()
        # 处理填空题
        for run in paragraph.runs:
            if run.underline:
                run.text = "______"
        text = paragraph.text.strip()
        # 处理括号
        text = re.sub(re.compile(r'（\s*[A-Z√×]\s*）'), '（  ）', text)
        # 处理大题
        if "答：" in text:
            answer += ans
            continue
        # 预处理完毕，匹配到题型
        if pattern_type.match(text):
            if question:
                question_dict[question] = answer
                question_list.append(question_dict)
            if question_type:
                bank_dict[question_type] = question_list
            question_type = text.split("、", 1)[-1]
            question_list = []
            question_dict = {}
            question = ""
            answer = ""
        else:
            if pattern_question.match(text):
                if question:
                    question_dict[question] = answer
                    question_list.append(question_dict)
                question_dict = {}
                question = ""
                answer = ""
                question += text + '\n'
                answer += ans + '\n'
            else:
                question += text + '\n'
                answer += ans + '\n'
    if question:
        question_dict[question] = answer
        question_list.append(question_dict)
    if question_type:
        bank_dict[question_type] = question_list

    return bank_dict


# 随机出题
def generate_random_problems(bank_dict):
    # 创建 Tkinter 应用程序
    root = tk.Tk()
    root.withdraw()
    type_cnt = 0
    doc_problem = Document()
    doc_answer = Document()
    for type, problem_dicts in bank_dict.items():
        str_type = chinese_numbers[type_cnt] + "、" + type
        doc_problem.add_paragraph(str_type)
        doc_answer.add_paragraph(str_type)
        type_cnt += 1
        problems_all = len(problem_dicts)
        # 输入你想生成的题数
        while True:
            try:
                problems_num = simpledialog.askstring \
                    (title=f"随机组卷系统",
                     prompt=f"{type}共{problems_all}题，请输入要生成{type}的题数：")
                if problems_num is None:
                    sys.exit()
                    # break
                problems_num = int(problems_num)
                if problems_num < 0:
                    raise ValueError
                break
            except (ValueError, TypeError):
                messagebox.showerror("提示", "请输入一个正确数字。")
        selected_problemdicts = random.sample(problem_dicts, min(problems_num, problems_all))
        problem_cnt = 1
        for problem in selected_problemdicts:
            text = next(iter(problem.keys()))  # 获取字典的第一个键
            answer = next(iter(problem.values()))  # 获取字典的第一个值n
            problem_withoutnum = re.sub(r'\d+\.', '', text)
            answer_withoutnum = re.sub(r'\d+\.', '', answer)
            text_problem = str(problem_cnt) + "." + problem_withoutnum
            text_answer = str(problem_cnt) + "." + answer_withoutnum
            # 打印题目
            doc_problem.add_paragraph(text_problem)
            # 打印答案
            doc_answer.add_paragraph(text_answer)
            problem_cnt += 1
    # 设置字体
    doc_problem.styles['Normal'].font.name = '宋体'
    doc_problem.styles['Normal'].font.size = Pt(12)
    doc_problem.styles['Normal'].font.unicode = True
    doc_problem.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc_answer.styles['Normal'].font.name = '宋体'
    doc_answer.styles['Normal'].font.size = Pt(12)
    doc_answer.styles['Normal'].font.unicode = True
    doc_answer.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc_problem.save('题目.docx')

    doc_answer.save('答案.docx')
    messagebox.showinfo("提示", "出卷完成！请在文件夹中查看题目和答案。")


bank_dict = generate_bank_dict()
generate_random_problems(bank_dict)
